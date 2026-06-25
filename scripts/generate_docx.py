import os
import docx
from docx.shared import Pt, Inches

def create_resume():
    doc = docx.Document()
    
    # Page setup - 0.5 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Name
    p_name = doc.add_paragraph()
    p_name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("Ayushkumar Burde")
    run_name.font.size = Pt(18)
    run_name.font.bold = True

    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    run_contact = p_contact.add_run(
        "ayushburde156@gmail.com | +91-8600820291 | Mumbai, India | "
        "LinkedIn: https://linkedin.com | GitHub: https://github.com | Website: https://website.com"
    )
    run_contact.font.size = Pt(9.5)

    # Helper function for sections
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        # Bottom border or division
        p_border = doc.add_paragraph()
        p_border.paragraph_format.space_after = Pt(4)
        run_border = p_border.add_run("―" * 60)
        run_border.font.size = Pt(6)
        run_border.font.color.rgb = docx.shared.RGBColor(211, 211, 211)

    # Summary
    add_section_header("Summary")
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_after = Pt(8)
    run_sum = p_sum.add_run(
        "Software Engineer Intern with Bachelor's in Computer Science and hands-on experience in software development, "
        "building AI-driven, cloud-based systems. Experienced in backend services, React, TypeScript, and Python, with "
        "familiarity in Agile development practices, version control, and modern technologies. Eager to contribute to "
        "projects and grow within innovative technical teams."
    )
    run_sum.font.size = Pt(10.5)

    # Experience
    add_section_header("Experience")
    p_job = doc.add_paragraph()
    run_job = p_job.add_run("CrystalTech Services Pvt Ltd — Software Engineer Intern")
    run_job.font.bold = True
    run_job.font.size = Pt(11)
    
    p_job_details = doc.add_paragraph()
    p_job_details.paragraph_format.space_after = Pt(3)
    run_jd = p_job_details.add_run("Jul 2024 - Dec 2024 | Indore, India")
    run_jd.font.italic = True
    run_jd.font.size = Pt(9.5)

    bullets = [
        "Developed microservices using Node.js and Express.js; deployed on AWS, enhancing cloud infrastructure performance by 40ms.",
        "Collaborated on React.js and TypeScript frontend; integrated RESTful APIs and followed Agile/SCRUM practices with version control via Git.",
        "Utilized MongoDB and Elasticsearch as NoSQL database solutions; supported CI/CD pipelines and debugging in dynamic systems."
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Education
    add_section_header("Education")
    p_edu1 = doc.add_paragraph()
    run_edu1 = p_edu1.add_run("Tulsiramji Gaikwad Patil College of Engineering and Technology")
    run_edu1.font.bold = True
    run_edu1.font.size = Pt(11)
    p_edu1_details = doc.add_paragraph()
    p_edu1_details.paragraph_format.space_after = Pt(4)
    run_ed1 = p_edu1_details.add_run("Master of Computer Applications | Graduation: 2025 | Nagpur, India")
    run_ed1.font.italic = True
    run_ed1.font.size = Pt(9.5)

    p_edu2 = doc.add_paragraph()
    run_edu2 = p_edu2.add_run("City Premier College")
    run_edu2.font.bold = True
    run_edu2.font.size = Pt(11)
    p_edu2_details = doc.add_paragraph()
    p_edu2_details.paragraph_format.space_after = Pt(4)
    run_ed2 = p_edu2_details.add_run("Bachelor of Computer Applications | Graduation: 2022 | Nagpur, India")
    run_ed2.font.italic = True
    run_ed2.font.size = Pt(9.5)

    # Skills
    add_section_header("Skills")
    skills = [
        ("Languages & Frameworks", "Python, React, TypeScript, RESTful APIs, NoSQL database, Elasticsearch, microservices architecture, cloud infrastructure, AI-driven, JavaScript, SQL, OOP, backend development, Agile development, version control, software development"),
        ("Systems & Tools", "Git, GitHub, Apache, CI/CD, Django, Node.js, Express.js, AWS, Azure, Google Cloud, Microsoft"),
        ("Leadership", "Problem-solving, Team collaboration, Agile/SCRUM, Collaborative, Understanding of requirements, Familiarity with development practices, Adaptability to new technologies")
    ]
    for category, details in skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(category + ": ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(details)
        r2.font.size = Pt(10)

    # Projects
    add_section_header("Projects")
    
    # Project 1
    p_p1 = doc.add_paragraph()
    r_p1 = p_p1.add_run("ResumeAI — React, FastAPI, Python, WeasyPrint, AI")
    r_p1.font.bold = True
    r_p1.font.size = Pt(11)
    doc.add_paragraph("Built an AI-powered ATS Resume Builder with LaTeX-based PDF generation, dynamic template rendering, and job description keyword optimization.", style='List Bullet')
    doc.add_paragraph("Engineered a robust PDF export engine using WeasyPrint and Jinja2, ensuring pixel-perfect layout preservation.", style='List Bullet')

    # Project 2
    p_p2 = doc.add_paragraph()
    r_p2 = p_p2.add_run("NotesApp — React.js, Node.js, Express.js, MongoDB, RESTful APIs, JWT, AWS, Git")
    r_p2.font.bold = True
    r_p2.font.size = Pt(11)
    doc.add_paragraph("Built full-stack app using React, Node.js, and MongoDB; deployed on AWS, serving 2K+ users with secure version control and development workflows.", style='List Bullet')

    # Project 3
    p_p3 = doc.add_paragraph()
    r_p3 = p_p3.add_run("CV Generator — Python, Django, SQLite, HTML5, CSS3, ReportLab, Google Cloud")
    r_p3.font.bold = True
    r_p3.font.size = Pt(11)
    doc.add_paragraph("Developed Python/Django app with PDF export; used Google tools and cloud storage for file handling, showcasing backend and technical control.", style='List Bullet')

    # Certifications
    add_section_header("Certifications")
    p_c1 = doc.add_paragraph("Complete 2025 Python Bootcamp From Scratch • Udemy (2024)", style='List Bullet')
    p_c1.paragraph_format.space_after = Pt(2)
    p_c2 = doc.add_paragraph("Django Masterclass • Udemy (2024)", style='List Bullet')
    p_c2.paragraph_format.space_after = Pt(2)

    doc.save("ayush_resume.docx")
    print("ayush_resume.docx created successfully!")

if __name__ == "__main__":
    create_resume()
