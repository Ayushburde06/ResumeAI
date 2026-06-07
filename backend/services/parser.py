import io
import pdfplumber
import fitz  # PyMuPDF
from docx import Document


# ── Magic-byte MIME validation ────────────────────────────────────────────────
# Validate actual file content regardless of the supplied filename/extension.
_PDF_MAGIC  = b"%PDF"
_ZIP_MAGIC  = b"PK\x03\x04"   # DOCX, XLSX, PPTX are all ZIP-based

def _check_mime(filename: str, file_bytes: bytes) -> None:
    """Raise ValueError if the file bytes do not match the expected type."""
    header = file_bytes[:8]
    lower  = filename.lower()

    if lower.endswith(".pdf"):
        if not header.startswith(_PDF_MAGIC):
            raise ValueError("The uploaded file does not appear to be a valid PDF.")
    elif lower.endswith(".docx"):
        if not header.startswith(_ZIP_MAGIC):
            raise ValueError("The uploaded file does not appear to be a valid DOCX.")


def _extract_pdf_annotation_links(file_bytes: bytes) -> list[str]:
    """Return all URI hyperlinks stored as PDF link annotations."""
    urls: list[str] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            for link in page.get_links():
                if link.get("kind") == fitz.LINK_URI:
                    uri = link.get("uri", "").strip()
                    if uri and uri.startswith(("http", "linkedin", "github", "mailto")):
                        urls.append(uri)
        doc.close()
    except Exception:
        pass
    return urls


def _extract_docx_hyperlinks(file_bytes: bytes) -> list[str]:
    """Return all hyperlink targets stored in DOCX relationships."""
    urls: list[str] = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                target = str(rel._target).strip()
                if target and not target.startswith("#"):
                    urls.append(target)
        # Also check headers, footers, and text boxes
        for section in doc.sections:
            for part in (section.header, section.footer):
                try:
                    for rel in part.part.rels.values():
                        if "hyperlink" in rel.reltype:
                            target = str(rel._target).strip()
                            if target and not target.startswith("#"):
                                urls.append(target)
                except Exception:
                    pass
    except Exception:
        pass
    return list(dict.fromkeys(urls))  # deduplicate, preserve order


def _append_links(base_text: str, urls: list[str]) -> str:
    """Append URLs that are not already present in the extracted text."""
    if not urls:
        return base_text
    new_urls = [u for u in urls if u not in base_text]
    if not new_urls:
        return base_text
    return base_text + "\n\nHyperlinks:\n" + "\n".join(new_urls)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    base_text = "\n".join(text_parts).strip()
    urls = _extract_pdf_annotation_links(file_bytes)
    return _append_links(base_text, urls)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    base_text = "\n".join(paragraphs).strip()
    urls = _extract_docx_hyperlinks(file_bytes)
    return _append_links(base_text, urls)


def parse_resume(filename: str, file_bytes: bytes) -> str:
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        _check_mime(filename, file_bytes)
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        _check_mime(filename, file_bytes)
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
    else:
        raise ValueError(f"Unsupported file type. Please upload a PDF or DOCX file.")
