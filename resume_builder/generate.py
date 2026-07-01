#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys, io
# Force UTF-8 output on Windows so status symbols print correctly
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

"""
generate.py -- Agentic Resume Generation Pipeline
=================================================
Closed loop: Validate JSON → Build DOCX → Render HTML → PDF via Playwright
           → Rasterize → Inspect pages → Auto-correct → Re-verify → Report

Usage:
    python generate.py                        # use resume_data.json
    python generate.py --data test_data/test_a_minimal.json
    python generate.py --stress-test          # run all 3 stress tests
    python generate.py --max-iter 8           # increase correction iterations
"""

import argparse
import asyncio
import json
import re
import sys
import textwrap
from dataclasses import dataclass, field
from pathlib import Path

# ── third-party guards ────────────────────────────────────────────────────────
def _require(module, install_hint):
    try:
        return __import__(module)
    except ImportError:
        print(f"ERROR: {install_hint}")
        sys.exit(1)

jsonschema  = _require("jsonschema",  "pip install jsonschema")
docx_mod    = _require("docx",        "pip install python-docx")
pdfium      = _require("pypdfium2",   "pip install pypdfium2")
PIL_mod     = _require("PIL",         "pip install Pillow")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

from playwright.async_api import async_playwright

# ── paths ─────────────────────────────────────────────────────────────────────
ROOT       = Path(__file__).parent
SCHEMA     = ROOT / "schema.json"
OUTPUT_DIR = ROOT / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

# ── typography params ─────────────────────────────────────────────────────────
@dataclass
class TypographyParams:
    font_size:          float = 10.5   # pt — body text
    name_font_size:     float = 20.0   # pt — header name
    section_font_size:  float = 11.0   # pt — section titles
    line_spacing:       float = 1.12   # multiple
    section_before:     float = 9.0    # pt — space before each section
    section_after:      float = 2.0    # pt — space after section title
    entry_before:       float = 4.0    # pt — space before each job/project header
    entry_after:        float = 0.0    # pt — space after header row
    bullet_before:      float = 1.0    # pt — space before each bullet
    contact_size:       float = 9.0    # pt — contact line font size
    margin_top:         float = 0.50   # inches
    margin_bottom:      float = 0.50   # inches
    margin_side:        float = 0.60   # inches

    def content_width_inches(self) -> float:
        """Usable content width (Letter page minus both side margins)."""
        return 8.5 - 2 * self.margin_side

    def right_tab_pos(self) -> float:
        """Position (in inches) for a right-aligned tab stop within content area."""
        return self.content_width_inches()

    def step_down(self) -> bool:
        """Reduce font/spacing by one notch. Returns False if already at minimum."""
        if self.font_size <= 9.5:
            return False
        self.font_size          = round(self.font_size - 0.25, 2)
        self.section_before     = max(4.0, self.section_before - 2.0)
        self.entry_before       = max(1.5, self.entry_before - 1.0)
        self.bullet_before      = max(0.0, self.bullet_before - 0.5)
        self.line_spacing       = max(1.0,  self.line_spacing - 0.04)
        return True

    def step_up(self) -> bool:
        """Increase font/spacing by one notch. Returns False if already at maximum."""
        if self.font_size >= 11.5:
            return False
        self.font_size          = round(self.font_size + 0.25, 2)
        self.section_before     = min(14.0, self.section_before + 2.0)
        self.entry_before       = min(7.0,  self.entry_before + 1.0)
        self.bullet_before      = min(3.0,  self.bullet_before + 0.5)
        self.line_spacing       = min(1.25, self.line_spacing + 0.04)
        return True


# ══════════════════════════════════════════════════════════════════════════════
# 1 ─ DATA VALIDATION
# ══════════════════════════════════════════════════════════════════════════════

def load_and_validate(data_path: Path) -> dict:
    """Load JSON and validate against schema. Raises on any violation."""
    if not data_path.exists():
        raise FileNotFoundError(f"Data file not found: {data_path}")

    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    with open(SCHEMA, encoding="utf-8") as f:
        schema = json.load(f)

    validator = jsonschema.Draft7Validator(schema)
    errors = sorted(validator.iter_errors(data), key=lambda e: e.path)
    if errors:
        print("\n── VALIDATION ERRORS ─────────────────────────────────────────")
        for e in errors:
            path = " → ".join(str(p) for p in e.path) or "(root)"
            print(f"  [{path}] {e.message}")
        raise ValueError(f"{len(errors)} validation error(s) in {data_path.name}")

    # Extra semantic checks
    _semantic_checks(data)

    print(f"  OK JSON valid ({data_path.name})")
    return data


def _semantic_checks(data: dict):
    """Catch truncated sentences, empty strings, suspiciously short text."""
    issues = []

    # Summary should end with proper punctuation
    summary = data.get("summary", "")
    if summary and not summary.rstrip().endswith(('.', '!', '?')):
        issues.append(f"summary does not end with proper punctuation: '…{summary[-30:]}'")

    # Each bullet should end with punctuation
    for job in data.get("experience", []):
        for i, b in enumerate(job.get("bullets", [])):
            if b and not b.rstrip().endswith(('.', '!', '?')):
                issues.append(
                    f"experience[{job['company']}] bullet {i+1} has no end punctuation: '…{b[-30:]}'"
                )
            if len(b.split()) < 5:
                issues.append(f"experience bullet looks truncated (< 5 words): '{b}'")

    for proj in data.get("projects", []):
        for i, b in enumerate(proj.get("bullets", [])):
            if b and not b.rstrip().endswith(('.', '!', '?')):
                issues.append(
                    f"project[{proj['name']}] bullet {i+1} has no end punctuation: '…{b[-30:]}'"
                )
            if len(b.split()) < 5:
                issues.append(f"project bullet looks truncated (< 5 words): '{b}'")

    if issues:
        print("\n── SEMANTIC WARNINGS ─────────────────────────────────────────")
        for w in issues:
            print(f"  WARN  {w}")
        # Warnings don't stop generation; only schema errors do


# ══════════════════════════════════════════════════════════════════════════════
# 2 ─ DOCX BUILDER
# ══════════════════════════════════════════════════════════════════════════════

# ── XML helpers ───────────────────────────────────────────────────────────────

def _add_bottom_border(para, color: str = "BBBBBB", sz: int = 8):
    """Add a thin bottom border to a paragraph (horizontal rule effect)."""
    p    = para._p
    pPr  = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))    # eighths of a point
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pBdr_existing = pPr.find(qn("w:pBdr"))
    if pBdr_existing is not None:
        pPr.remove(pBdr_existing)
    pPr.append(pBdr)


def _add_right_tab(para, inches: float):
    """Add a right-aligned tab stop at `inches` from left margin."""
    pPr  = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(inches * 1440)))   # twips
    tabs.append(tab)


def _set_spacing(para, before: float, after: float, line_mult: float):
    pf = para.paragraph_format
    pf.space_before     = Pt(before)
    pf.space_after      = Pt(after)
    pf.line_spacing     = line_mult
    # Keep paragraph together with next where possible
    pf.keep_with_next   = True


def _font(run, size_pt: float, bold=False, italic=False,
          color: RGBColor | None = None):
    run.bold       = bold
    run.italic     = italic
    run.font.size  = Pt(size_pt)
    if color:
        run.font.color.rgb = color


GREY = RGBColor(0x77, 0x77, 0x77)


# ── section header ─────────────────────────────────────────────────────────────

def _section_header(doc: Document, text: str, p: TypographyParams):
    para = doc.add_paragraph()
    _set_spacing(para, p.section_before, p.section_after, p.line_spacing)
    run = para.add_run(text.upper())
    _font(run, p.section_font_size, bold=True)
    _add_bottom_border(para)
    return para


# ── two-column entry header (company/project | date) ─────────────────────────

def _entry_row(doc: Document, left_text: str, right_text: str,
               p: TypographyParams, bold_left=True, italic_left=False):
    para = doc.add_paragraph()
    _set_spacing(para, p.entry_before, p.entry_after, p.line_spacing)
    _add_right_tab(para, p.right_tab_pos())
    run_l = para.add_run(left_text)
    _font(run_l, p.font_size, bold=bold_left, italic=italic_left)
    para.add_run("\t")
    run_r = para.add_run(right_text)
    _font(run_r, p.contact_size, color=GREY)
    return para


def _sub_row(doc: Document, left_text: str, right_text: str,
             p: TypographyParams):
    para = doc.add_paragraph()
    _set_spacing(para, 0, 0, p.line_spacing)
    _add_right_tab(para, p.right_tab_pos())
    run_l = para.add_run(left_text)
    _font(run_l, p.font_size, italic=True)
    para.add_run("\t")
    run_r = para.add_run(right_text)
    _font(run_r, p.contact_size, color=GREY)
    return para


# ── bullet ────────────────────────────────────────────────────────────────────

def _bullet(doc: Document, text: str, p: TypographyParams):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before    = Pt(p.bullet_before)
    para.paragraph_format.space_after     = Pt(0)
    para.paragraph_format.line_spacing    = p.line_spacing
    para.paragraph_format.left_indent     = Inches(0.20)
    para.paragraph_format.first_line_indent = Inches(-0.15)
    para.paragraph_format.keep_with_next  = False
    run = para.add_run(text)
    _font(run, p.font_size)
    return para


# ── main docx builder ─────────────────────────────────────────────────────────

def build_docx(data: dict, p: TypographyParams, out_path: Path) -> Path:
    doc = Document()

    # ── page setup (Letter) ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width      = Inches(8.5)
    section.page_height     = Inches(11.0)
    section.left_margin     = Inches(p.margin_side)
    section.right_margin    = Inches(p.margin_side)
    section.top_margin      = Inches(p.margin_top)
    section.bottom_margin   = Inches(p.margin_bottom)

    # ── default paragraph style ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name    = "Calibri"
    style.font.size    = Pt(p.font_size)
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── HEADER: name ─────────────────────────────────────────────────────────
    contact = data["contact"]
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(name_para, 0, 2, 1.0)
    run = name_para.add_run(contact["name"])
    _font(run, p.name_font_size, bold=True)

    # ── HEADER: contact info ─────────────────────────────────────────────────
    parts = [
        contact.get("email", ""),
        contact.get("phone", ""),
        contact.get("location", ""),
    ]
    for key in ("linkedin", "github", "website"):
        val = contact.get(key, "").strip()
        if val:
            # show domain only (cleaner in DOCX)
            label = key.capitalize()
            parts.append(f"{label}: {val.replace('https://', '').replace('http://', '')}")

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(contact_para, 0, 4, 1.0)
    run = contact_para.add_run(" · ".join(p for p in parts if p))
    _font(run, p.contact_size, color=GREY)

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    _section_header(doc, "Summary", p)
    sum_para = doc.add_paragraph()
    _set_spacing(sum_para, 0, 0, p.line_spacing)
    run = sum_para.add_run(data["summary"])
    _font(run, p.font_size)

    # ── EXPERIENCE ───────────────────────────────────────────────────────────
    if data.get("experience"):
        _section_header(doc, "Experience", p)
        for job in data["experience"]:
            _entry_row(doc, job["company"], job.get("dates", ""), p, bold_left=True)
            _sub_row(doc, job["role"], job.get("location", ""), p)
            for b in job.get("bullets", []):
                _bullet(doc, b, p)

    # ── EDUCATION ────────────────────────────────────────────────────────────
    if data.get("education"):
        _section_header(doc, "Education", p)
        for edu in data["education"]:
            _entry_row(doc, edu["institution"], edu.get("year", ""), p, bold_left=True)
            _sub_row(doc, edu["degree"], edu.get("location", ""), p)

    # ── SKILLS ───────────────────────────────────────────────────────────────
    if data.get("skills"):
        _section_header(doc, "Skills", p)
        for cat, items in data["skills"].items():
            sp = doc.add_paragraph()
            _set_spacing(sp, 1, 0, p.line_spacing)
            run_cat = sp.add_run(f"{cat}: ")
            _font(run_cat, p.font_size, bold=True)
            run_items = sp.add_run(", ".join(items))
            _font(run_items, p.font_size)

    # ── PROJECTS ─────────────────────────────────────────────────────────────
    if data.get("projects"):
        _section_header(doc, "Projects", p)
        for proj in data["projects"]:
            proj_label = proj["name"]
            if proj.get("tech"):
                proj_label += f"  |  {proj['tech']}"
            _entry_row(doc, proj_label, proj.get("dates", ""), p, bold_left=True)
            for b in proj.get("bullets", []):
                _bullet(doc, b, p)

    # ── CERTIFICATIONS ───────────────────────────────────────────────────────
    if data.get("certifications"):
        _section_header(doc, "Certifications", p)
        for cert in data["certifications"]:
            cp = doc.add_paragraph()
            _set_spacing(cp, 1, 0, p.line_spacing)
            run_name = cp.add_run(cert["name"])
            _font(run_name, p.font_size, bold=True)
            run_rest = cp.add_run(f"  ·  {cert['issuer']}  ·  {cert['year']}")
            _font(run_rest, p.font_size, color=GREY)

    doc.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# 3 ─ HTML BUILDER  (verification proxy — same data, same structure)
# ══════════════════════════════════════════════════════════════════════════════

def build_html(data: dict, p: TypographyParams) -> str:
    """Generate an HTML page that mirrors the DOCX layout for pixel verification."""

    def esc(text: str) -> str:
        return (text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace('"', "&quot;"))

    sections_html = []

    def section(title: str, body: str):
        sections_html.append(
            f'<div class="section">'
            f'<div class="section-title">{esc(title.upper())}</div>'
            f'{body}'
            f'</div>'
        )

    # SUMMARY
    section("Summary", f'<p class="summary">{esc(data["summary"])}</p>')

    # EXPERIENCE
    if data.get("experience"):
        exp_html = ""
        for job in data["experience"]:
            exp_html += f'''
<div class="entry">
  <div class="entry-header">
    <span class="entry-main">{esc(job["company"])}</span>
    <span class="entry-meta">{esc(job.get("dates",""))}</span>
  </div>
  <div class="entry-header">
    <span class="entry-sub">{esc(job["role"])}</span>
    <span class="entry-meta">{esc(job.get("location",""))}</span>
  </div>
  <ul>{"".join(f"<li>{esc(b)}</li>" for b in job.get("bullets",[]))}</ul>
</div>'''
        section("Experience", exp_html)

    # EDUCATION
    if data.get("education"):
        edu_html = ""
        for edu in data["education"]:
            edu_html += f'''
<div class="entry">
  <div class="entry-header">
    <span class="entry-main">{esc(edu["institution"])}</span>
    <span class="entry-meta">{esc(edu.get("year",""))}</span>
  </div>
  <div class="entry-header">
    <span class="entry-sub">{esc(edu["degree"])}</span>
    <span class="entry-meta">{esc(edu.get("location",""))}</span>
  </div>
</div>'''
        section("Education", edu_html)

    # SKILLS
    if data.get("skills"):
        sk_html = '<div class="skills">'
        for cat, items in data["skills"].items():
            sk_html += (
                f'<div class="skill-row">'
                f'<strong>{esc(cat)}:</strong> {esc(", ".join(items))}'
                f'</div>'
            )
        sk_html += '</div>'
        section("Skills", sk_html)

    # PROJECTS
    if data.get("projects"):
        proj_html = ""
        for proj in data["projects"]:
            label = proj["name"]
            if proj.get("tech"):
                label += f'  |  {proj["tech"]}'
            proj_html += f'''
<div class="entry">
  <div class="entry-header">
    <span class="entry-main">{esc(label)}</span>
    <span class="entry-meta">{esc(proj.get("dates",""))}</span>
  </div>
  <ul>{"".join(f"<li>{esc(b)}</li>" for b in proj.get("bullets",[]))}</ul>
</div>'''
        section("Projects", proj_html)

    # CERTIFICATIONS
    if data.get("certifications"):
        cert_html = '<div class="certs">'
        for cert in data["certifications"]:
            cert_html += (
                f'<div class="cert-row">'
                f'<strong>{esc(cert["name"])}</strong>'
                f' &nbsp;·&nbsp; {esc(cert["issuer"])}'
                f' &nbsp;·&nbsp; {esc(cert["year"])}'
                f'</div>'
            )
        cert_html += '</div>'
        section("Certifications", cert_html)

    # Contact
    contact = data["contact"]
    contact_parts = [
        contact.get("email",""),
        contact.get("phone",""),
        contact.get("location",""),
    ]
    for key in ("linkedin","github","website"):
        v = contact.get(key,"").strip()
        if v:
            contact_parts.append(v.replace("https://","").replace("http://",""))
    contact_str = " &nbsp;·&nbsp; ".join(esc(x) for x in contact_parts if x)

    css = f"""
    @page {{ size: letter; margin: {p.margin_top}in {p.margin_side}in {p.margin_bottom}in; }}
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
        font-family: 'Calibri', 'Arial', sans-serif;
        font-size: {p.font_size}pt;
        line-height: {p.line_spacing};
        color: #141414;
        -webkit-print-color-adjust: exact;
        print-color-adjust: exact;
    }}
    .name {{
        font-size: {p.name_font_size}pt;
        font-weight: bold;
        text-align: center;
        margin-bottom: 3pt;
    }}
    .contact {{
        font-size: {p.contact_size}pt;
        text-align: center;
        color: #777;
        margin-bottom: 6pt;
        word-spacing: 2px;
    }}
    .section {{
        margin-top: {p.section_before}pt;
    }}
    .section-title {{
        font-size: {p.section_font_size}pt;
        font-weight: bold;
        letter-spacing: 0.04em;
        border-bottom: 1.2px solid #BBBBBB;
        padding-bottom: 1.5px;
        margin-bottom: {p.section_after}pt;
    }}
    .summary {{ margin-top: 1pt; }}
    .entry {{
        margin-top: {p.entry_before}pt;
        page-break-inside: avoid;
    }}
    .entry-header {{
        display: flex;
        justify-content: space-between;
        align-items: baseline;
        gap: 8px;
    }}
    .entry-main {{ font-weight: bold; }}
    .entry-sub  {{ font-style: italic; }}
    .entry-meta {{ font-size: {p.contact_size}pt; color: #777; white-space: nowrap; flex-shrink: 0; }}
    ul {{
        margin: 0;
        padding-left: 16pt;
        list-style-type: disc;
    }}
    li {{
        margin-top: {p.bullet_before}pt;
        margin-bottom: 0;
    }}
    .skills, .certs {{ margin-top: 1pt; }}
    .skill-row {{ margin-top: 1.5pt; }}
    .cert-row  {{ margin-top: 1.5pt; }}
    """

    body = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>{css}</style>
</head>
<body>
<div class="name">{esc(contact["name"])}</div>
<div class="contact">{contact_str}</div>
{"".join(sections_html)}
</body>
</html>"""

    return body


# ══════════════════════════════════════════════════════════════════════════════
# 4 ─ PLAYWRIGHT RENDERER  (HTML → PDF)
# ══════════════════════════════════════════════════════════════════════════════

async def _render_pdf(html: str, out_path: Path) -> Path:
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=True)
        page    = await browser.new_page()
        await page.set_content(html, wait_until="networkidle")
        await page.pdf(
            path             = str(out_path),
            format           = "Letter",
            print_background = True,
            prefer_css_page_size = True,
        )
        await browser.close()
    return out_path


def render_pdf(html: str, out_path: Path) -> Path:
    """Synchronous wrapper around the async Playwright PDF renderer."""
    loop = asyncio.ProactorEventLoop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(_render_pdf(html, out_path))
    finally:
        loop.close()


# ══════════════════════════════════════════════════════════════════════════════
# 5 ─ RASTERIZER  (PDF → images via pypdfium2)
# ══════════════════════════════════════════════════════════════════════════════

def rasterize(pdf_path: Path, out_dir: Path, dpi: int = 150) -> list[Path]:
    """Convert every page of the PDF to a PNG image. Returns list of image paths."""
    doc     = pdfium.PdfDocument(str(pdf_path))
    scale   = dpi / 72.0
    images  = []
    for i, page in enumerate(doc):
        bitmap   = page.render(scale=scale, rotation=0)
        pil_img  = bitmap.to_pil()
        img_path = out_dir / f"page_{i+1:02d}.png"
        pil_img.save(img_path, "PNG")
        images.append(img_path)
        page.close()
    doc.close()
    return images


# ══════════════════════════════════════════════════════════════════════════════
# 6 ─ VISUAL INSPECTOR  (Pillow-based pixel analysis)
# ══════════════════════════════════════════════════════════════════════════════

def _fill_ratio(img: Image.Image, region=None) -> float:
    """
    Return fraction of pixels in `region` that are NOT near-white.
    A white/near-white pixel has all channels >= 230.
    """
    if region:
        img = img.crop(region)
    rgb = img.convert("RGB")
    pixels = list(rgb.getdata())
    non_white = sum(1 for r, g, b in pixels if min(r, g, b) < 230)
    return non_white / len(pixels) if pixels else 0.0


def analyze_images(image_paths: list[Path]) -> list[str]:
    """
    Inspect rendered page images and return a list of issue codes:
      "overflow"            - text too close to bottom margin on page 1
      "sparse_bottom"       - bottom 35% of page 1 is nearly empty (content could expand)
      "second_page_sparse"  - last page has < 18% fill (content should reflow to page 1)
    """
    issues: list[str] = []

    img1 = Image.open(image_paths[0])
    w, h = img1.size

    # Zone dimensions (in pixels relative to 150 dpi rendering)
    bottom_margin_zone = (0, int(h * 0.93), w, h)        # last 7%  — overflow guard
    bottom_third_zone  = (0, int(h * 0.65), w, h)        # bottom 35% — sparseness check

    fill_margin = _fill_ratio(img1, bottom_margin_zone)
    fill_bottom = _fill_ratio(img1, bottom_third_zone)

    # Text bleeding into the very bottom margin
    if fill_margin > 0.04:
        issues.append("overflow")

    # Bottom third of page 1 nearly empty (and only 1 page) → expand
    if len(image_paths) == 1 and fill_bottom < 0.06:
        issues.append("sparse_bottom")

    # If a second page exists but is mostly blank → compress to fix reflow
    if len(image_paths) >= 2:
        last_img  = Image.open(image_paths[-1])
        fill_last = _fill_ratio(last_img)
        if fill_last < 0.18:
            issues.append("second_page_sparse")

    return issues


# ══════════════════════════════════════════════════════════════════════════════
# 7 ─ MAIN SELF-CORRECTION LOOP
# ══════════════════════════════════════════════════════════════════════════════

def _banner(msg: str):
    bar = "─" * 60
    print(f"\n{bar}\n  {msg}\n{bar}")


def run_pipeline(data: dict, label: str, max_iter: int = 6) -> dict:
    """
    Full generate→verify→fix loop for one dataset.
    Returns a result dict with final image paths and params.
    """
    _banner(f"Generating: {label}")

    stem     = label.replace(" ", "_").lower()
    docx_out = OUTPUT_DIR / f"{stem}.docx"
    pdf_out  = OUTPUT_DIR / f"{stem}.pdf"
    img_dir  = OUTPUT_DIR / stem
    img_dir.mkdir(exist_ok=True)

    params = TypographyParams()
    history: list[dict] = []

    for iteration in range(1, max_iter + 1):
        print(f"\n  Iteration {iteration}  |  font={params.font_size}pt  "
              f"line={params.line_spacing:.2f}  sec_before={params.section_before}pt")

        # Step 1: Build DOCX
        build_docx(data, params, docx_out)
        print(f"    OK DOCX written -> {docx_out.name}")

        # Step 2: Build HTML + render PDF
        html = build_html(data, params)
        render_pdf(html, pdf_out)
        print(f"    OK PDF rendered  -> {pdf_out.name}")

        # Step 3: Rasterize
        # Clear old images
        for f in img_dir.glob("page_*.png"):
            f.unlink()
        images = rasterize(pdf_out, img_dir)
        print(f"    OK Rasterized    -> {len(images)} page(s)")

        # Step 4: Inspect
        issues = analyze_images(images)
        history.append({"iteration": iteration, "issues": issues[:], "params": {
            "font_size": params.font_size,
            "line_spacing": params.line_spacing,
        }})

        if not issues:
            print(f"    PASS Visual check PASSED -- no layout issues detected")
            break

        print(f"    FAIL Issues found: {issues}")

        # Step 5: Identify root cause and adjust
        if "overflow" in issues:
            print("      Root cause: font/spacing too large → content overflows margin")
            if not params.step_down():
                print("      Already at minimum size — cannot reduce further")
                break

        elif "second_page_sparse" in issues:
            print("      Root cause: section/bullet spacing spreads content to a sparse 2nd page")
            if not params.step_down():
                print("      Already at minimum size — cannot reduce further")
                break

        elif "sparse_bottom" in issues:
            print("      Root cause: font/spacing too small → page has large unused whitespace")
            if not params.step_up():
                print("      Already at maximum size — cannot increase further")
                break

    else:
        print(f"  ⚠  Reached max iterations ({max_iter}) without a clean pass")

    print(f"\n  Final output:")
    print(f"    DOCX  → {docx_out}")
    print(f"    PDF   → {pdf_out}")
    print(f"    Pages → {', '.join(p.name for p in images)}")

    return {
        "label":   label,
        "docx":    docx_out,
        "pdf":     pdf_out,
        "images":  images,
        "params":  params,
        "history": history,
        "passed":  len(history) > 0 and not history[-1]["issues"],
    }


# ══════════════════════════════════════════════════════════════════════════════
# 8 ─ STRESS TESTS
# ══════════════════════════════════════════════════════════════════════════════

def run_stress_tests(max_iter: int):
    """Run Tests A (minimal), B (baseline), C (heavy) and report results."""
    test_files = {
        "Test A — Minimal":  ROOT / "test_data" / "test_a_minimal.json",
        "Test B — Baseline": ROOT / "resume_data.json",
        "Test C — Heavy":    ROOT / "test_data" / "test_c_heavy.json",
    }

    results = []
    for label, path in test_files.items():
        if not path.exists():
            print(f"  SKIP: {path} not found")
            continue
        data = load_and_validate(path)
        r = run_pipeline(data, label, max_iter=max_iter)
        results.append(r)

    _banner("STRESS TEST REPORT")
    for r in results:
        status  = "PASS" if r["passed"] else "NEEDS REVIEW"
        iters   = len(r["history"])
        pages   = len(r["images"])
        print(f"\n  {r['label']}")
        print(f"    Status     : {status}")
        print(f"    Pages      : {pages}")
        print(f"    Iterations : {iters}")
        print(f"    Final font : {r['params'].font_size}pt")
        print(f"    DOCX       : {r['docx'].name}")
        print(f"    Screenshots: {r['images'][0].parent.name}/page_*.png")

    print()
    return results


# ══════════════════════════════════════════════════════════════════════════════
# 9 ─ ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(
        description="Agentic resume generation pipeline"
    )
    ap.add_argument("--data",        default="resume_data.json",
                    help="Path to resume_data.json (default: resume_data.json)")
    ap.add_argument("--stress-test", action="store_true",
                    help="Run all 3 stress tests (A/B/C)")
    ap.add_argument("--max-iter",    type=int, default=6,
                    help="Max self-correction iterations (default: 6)")
    args = ap.parse_args()

    if args.stress_test:
        run_stress_tests(max_iter=args.max_iter)
        return

    data_path = ROOT / args.data if not Path(args.data).is_absolute() else Path(args.data)
    data = load_and_validate(data_path)
    result = run_pipeline(data, label="resume", max_iter=args.max_iter)

    if result["passed"]:
        print("\n  Pipeline complete. Resume is verified and ready.")
    else:
        print("\n  Pipeline complete with warnings — review screenshots in output/")

    print(f"\n  Quick access:")
    print(f"    edit   : resume_builder/resume_data.json")
    print(f"    regen  : python resume_builder/generate.py")
    print(f"    output : resume_builder/output/resume.docx")


if __name__ == "__main__":
    main()
